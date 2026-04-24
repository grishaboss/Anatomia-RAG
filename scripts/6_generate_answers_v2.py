#!/usr/bin/env python3
"""
Шаг 6: Генерация ответов с разбивкой по авторам и межавторским сравнением.

Что делает:
  • Загружает обогащённые вопросы из data/enriched_questions.json
    (предварительно запустите: python scripts/5_enrich_questions.py)
  • Для каждого вопроса делает RAG-запросы отдельно по каждому автору:
      queries = [основной вопрос] + [подвопросы]
  • Передаёт LLM контекст каждого автора в отдельных блоках
  • LLM сравнивает авторов: совпадающие факты → в ответ без пометок,
    расхождения → явно ("Гайворонский выделяет X, тогда как Сапин описывает Y")
  • Сохраняет форматированный DOCX с чекпойнтами

Использование:
    python scripts/6_generate_answers_v2.py
    python scripts/6_generate_answers_v2.py --resume
    python scripts/6_generate_answers_v2.py --question-range 1-5
    python scripts/6_generate_answers_v2.py --no-llm
    python scripts/6_generate_answers_v2.py --backend openai --model gpt-4o
    python scripts/6_generate_answers_v2.py --enrichment-file data/processed/enriched_questions_from_chatgpt-5o.json
"""
from __future__ import annotations

import argparse
import json
import logging
import re
import sys
from pathlib import Path

import chromadb
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sentence_transformers import SentenceTransformer
from tqdm import tqdm

sys.path.insert(0, str(Path(__file__).parent))
from utils import ROOT, call_llm, format_context, get_device, load_config, retrieve

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger(__name__)


# ──────────────────────────────────────────────
#  Системный промпт
# ──────────────────────────────────────────────

SYSTEM_PROMPT_V2 = """Ты — редактор учебного конспекта. Твоя задача: выбрать и организовать информацию ИСКЛЮЧИТЕЛЬНО из предоставленных фрагментов учебников.

КАТЕГОРИЧЕСКИЕ ПРАВИЛА:
1. ЗАПРЕЩЕНО добавлять любые факты, цифры, термины или утверждения, которых НЕТ в предоставленных фрагментах
2. ЗАПРЕЩЕНО использовать свои знания по анатомии — только текст из фрагментов
3. Если по какому-то подвопросу информации в фрагментах нет — напиши: [нет в источниках]
4. Пиши максимально близко к оригинальному тексту учебника — минимально редактируй формулировки
5. Если Гайворонский и Сапин говорят разное — укажи: «По Гайворонскому: ..., по Сапину: ...»
6. Если оба автора говорят одно и то же — пиши один раз без атрибуции
7. НЕ придумывай структуру, которой нет в источниках; раскрывай только то, что есть в тексте"""


# ──────────────────────────────────────────────
#  Загрузка вопросов
# ──────────────────────────────────────────────

_SECTION_NAMES: dict[str, str] = {
    "I":   "I. Общие вопросы анатомии. Анатомия опорно-двигательного аппарата",
    "II":  "II. Спланхнология",
    "III": "III. Анатомия сердечно-сосудистой и лимфатической систем",
    "IV":  "IV. Нейроанатомия. Анатомия органов чувств.",
}


def _section_for_exam_num(exam_num: str) -> str:
    """'I.3' → 'I. Опорно-двигательный аппарат'."""
    roman = exam_num.split(".")[0]
    return _SECTION_NAMES.get(roman, f"Раздел {roman}")


def load_enriched_questions(path: Path) -> list[dict]:
    with open(path, encoding="utf-8") as f:
        data = json.load(f)

    # Поддержка формата ChatGPT-5o / LLM: {exam_num, exam_text, sub_questions}
    if data and "exam_num" in data[0]:
        normalized: list[dict] = []
        for i, item in enumerate(data, 1):
            normalized.append({
                "number":      i,                                   # порядковый для фильтрации
                "exam_num":    item["exam_num"],                    # 'I.3' для отображения
                "question":    item["exam_text"],
                "section":     _section_for_exam_num(item["exam_num"]),
                "sub_questions": item.get("sub_questions", []),
            })
        return normalized

    return data


# ──────────────────────────────────────────────
#  RAG по автору
# ──────────────────────────────────────────────

def retrieve_for_author(
    queries: list[str],
    author: str,
    collection,
    embed_model,
    config: dict,
    max_chunks: int,
    top_k_per_query: int,
    hyde: bool = False,
) -> list[dict]:
    """Собрать и дедуплицировать чанки для одного автора по списку запросов."""
    seen_texts: set[str] = set()
    all_chunks: list[dict] = []

    # Временный конфиг с уменьшенным top_k для одного sub-query
    override_cfg = {**config, "retrieval": {**config["retrieval"], "top_k": top_k_per_query}}

    for i, query in enumerate(queries):
        # HyDE применяем только к первому запросу (основной вопрос),
        # подвопросы уже достаточно конкретны.
        use_hyde = hyde and i == 0
        chunks = retrieve(query, collection, embed_model, override_cfg,
                          author_filter=author, hyde=use_hyde)
        for chunk in chunks:
            text = chunk["text"]
            if text not in seen_texts:
                seen_texts.add(text)
                all_chunks.append(chunk)

    # Отсортировать по score и взять top max_chunks
    all_chunks.sort(key=lambda c: c["score"], reverse=True)
    return all_chunks[:max_chunks]


# ──────────────────────────────────────────────
#  Формирование промпта
# ──────────────────────────────────────────────

def build_user_prompt(
    question: dict,
    ctx_by_author: dict[str, str],
) -> str:
    sub_qs = question.get("sub_questions", [])
    sub_list = "\n".join(f"  • {sq['text']}" for sq in sub_qs) if sub_qs else "—"

    context_blocks = []
    for author_name, ctx in ctx_by_author.items():
        block_text = ctx if ctx.strip() else "(нет данных в базе знаний)"
        context_blocks.append(f"=== {author_name} ===\n{block_text}")

    return (
        "# ФРАГМЕНТЫ ИЗ УЧЕБНИКОВ (используй ТОЛЬКО этот текст):\n\n"
        + "\n\n".join(context_blocks)
        + f"\n\n{'='*60}\n"
        f"# ВОПРОС: {question['question']}\n\n"
        f"# АСПЕКТЫ ДЛЯ ОХВАТА:\n{sub_list}\n\n"
        "# ЗАДАНИЕ:"
        "Напиши единый связный ответ на экзаменационный вопрос выше. "
        "Используй аспекты как ориентиры для полноты охвата — каждый из них должен быть "
        "раскрыт в тексте, но НЕ оформляй их как отдельные пункты с заголовками. "
        "Пиши кратко и по существу, близко к оригинальному тексту учебника - примерно для ответа на 5 минут. "
        "Если какой-то аспект не найден во фрагментах — просто пропусти его. "
        "НЕ добавляй ничего от себя."
    )


# ──────────────────────────────────────────────
#  Чекпойнт
# ──────────────────────────────────────────────

def ckpt_path(output_path: Path) -> Path:
    return output_path.with_suffix(".checkpoint.json")


def load_checkpoint(output_path: Path) -> dict[int, dict]:
    p = ckpt_path(output_path)
    if p.exists():
        with open(p, encoding="utf-8") as f:
            raw = json.load(f)
        return {int(k): v for k, v in raw.items()}
    return {}


def save_checkpoint(output_path: Path, data: dict[int, dict]) -> None:
    with open(ckpt_path(output_path), "w", encoding="utf-8") as f:
        json.dump({str(k): v for k, v in data.items()}, f, ensure_ascii=False, indent=2)


# ──────────────────────────────────────────────
#  Построение DOCX
# ──────────────────────────────────────────────

def _unique_sources_by_author(chunks_by_author: dict[str, list[dict]]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for author, chunks in chunks_by_author.items():
        for chunk in chunks:
            m = chunk["meta"]
            s = m.get("source_name", "")
            page = m.get("page")
            label = f"{author} — {s} (стр. {page})" if page else f"{author} — {s}"
            if label not in seen:
                seen.add(label)
                result.append(label)
    return result


def _collect_figures(chunks_by_author: dict[str, list[dict]], max_per_question: int = 6) -> list[Path]:
    """Собрать пути к существующим фигурам из метаданных чанков (без дублей)."""
    seen: set[str] = set()
    figures: list[Path] = []
    for chunks in chunks_by_author.values():
        for chunk in chunks:
            raw = chunk["meta"].get("figures", "[]")
            try:
                paths = json.loads(raw) if isinstance(raw, str) else (raw or [])
            except (json.JSONDecodeError, TypeError):
                paths = []
            for p in paths:
                if p not in seen:
                    seen.add(p)
                    full = ROOT / p
                    if full.exists():
                        figures.append(full)
                    if len(figures) >= max_per_question:
                        return figures
    return figures


_IMAGE_TAG_RE_CLEAN = re.compile(r'<!--\s*image\s*-->', re.IGNORECASE)


def build_output_doc(questions: list[dict], results: list[dict]) -> DocxDocument:
    doc = DocxDocument()

    title = doc.add_heading(
        "Ответы на экзаменационные вопросы\nпо анатомии человека", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        "Версия с межавторским сравнением (Гайворонский / Сапин / Якимов)"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    current_section: str | None = None

    for question, res in zip(questions, results):
        section = question.get("section", "")
        if section != current_section:
            current_section = section
            if current_section:
                sec_heading = doc.add_heading(current_section, level=1)
                if sec_heading.runs:
                    sec_heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0x76)

        # Заголовок вопроса
        q_heading = doc.add_heading(
            f"Вопрос {question.get('exam_num', question['number'])}. {question['question']}", level=2
        )
        for run in q_heading.runs:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        # Подвопросы (мелко, курсивом)
        sub_qs = question.get("sub_questions", [])
        if sub_qs:
            sub_para = doc.add_paragraph()
            sub_run = sub_para.add_run(
                "Аспекты: " + "  •  ".join(sq["text"] for sq in sub_qs)
            )
            sub_run.italic = True
            sub_run.font.size = Pt(9)
            sub_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        # Ответ (очищаем от тегов <!-- image -->)
        answer_text = _IMAGE_TAG_RE_CLEAN.sub("", res.get("answer", ""))
        for line in answer_text.split("\n"):
            line = line.strip()
            if line:
                p = doc.add_paragraph(line)
                p.style = "Normal"
                p.paragraph_format.space_after = Pt(2)

        # Рисунки из извлечённых чанков
        figures = _collect_figures(res.get("chunks_by_author", {}))
        if figures:
            fig_para = doc.add_paragraph()
            fig_run = fig_para.add_run("Иллюстрации:")
            fig_run.bold = True
            fig_run.font.size = Pt(9)
            for fig_path in figures:
                try:
                    doc.add_picture(str(fig_path), width=Inches(5.5))
                    cap = doc.add_paragraph(fig_path.stem)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cap.runs:
                        run.italic = True
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                except Exception:
                    pass  # битый/нечитаемый файл — пропустить

        # Источники
        sources = _unique_sources_by_author(res.get("chunks_by_author", {}))
        if sources:
            src_para = doc.add_paragraph()
            src_para.paragraph_format.space_before = Pt(4)
            label_run = src_para.add_run("Источники: ")
            label_run.bold = True
            label_run.font.size = Pt(9)
            src_run = src_para.add_run(";  ".join(sources))
            src_run.italic = True
            src_run.font.size = Pt(9)
            src_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        doc.add_paragraph()

    return doc


# ──────────────────────────────────────────────
#  Точка входа
# ──────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Генерация ответов с межавторским сравнением (Гайворонский / Сапин / Якимов)"
    )
    parser.add_argument("--backend",        choices=["ollama", "openai"], help="LLM backend")
    parser.add_argument("--model",          help="Название модели")
    parser.add_argument("--no-llm",         action="store_true", help="Только контекст, без LLM")
    parser.add_argument("--resume",         action="store_true", help="Продолжить прерванную генерацию")
    parser.add_argument(
        "--question-range",
        help="Диапазон номеров вопросов, напр. 1-10",
        metavar="START-END",
    )
    parser.add_argument(
        "--enrichment-file",
        help="Путь к файлу обогащённых вопросов (переопределяет config.yaml)",
        metavar="PATH",
    )
    args = parser.parse_args()

    config  = load_config()
    device  = get_device(config)
    gen_cfg = config.get("generation_v2", {})

    enrichment_cfg = config.get("enrichment", {})
    llm_file    = enrichment_cfg.get("llm_file")
    output_file = enrichment_cfg.get("output_file", "data/enriched_questions.json")

    if args.enrichment_file:
        # CLI-аргумент имеет наивысший приоритет
        enriched_file = Path(args.enrichment_file)
        if not enriched_file.is_absolute():
            enriched_file = ROOT / enriched_file
        logger.info(f"Используем файл из аргумента: {enriched_file.name}")
    elif llm_file and (ROOT / llm_file).exists():
        # Предпочитаем LLM-обогащение (ChatGPT-5o / Claude), если файл задан и существует
        enriched_file = ROOT / llm_file
        logger.info(f"Используем LLM-обогащение: {enriched_file.name}")
    else:
        enriched_file = ROOT / output_file
        logger.info(f"Используем embedding-обогащение: {enriched_file.name}")

    output_path = ROOT / gen_cfg.get("answers_v2_file", "output/Ответы_v2.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    max_chunks      = int(gen_cfg.get("max_chunks_per_author", 10))
    top_k_per_query = int(gen_cfg.get("top_k_per_query",        4))
    hyde            = bool(config.get("retrieval", {}).get("hyde", False))

    if not enriched_file.exists():
        logger.error(f"Файл обогащённых вопросов не найден: {enriched_file}")
        logger.error("Сначала запустите:  python scripts/5_enrich_questions.py")
        sys.exit(1)

    questions = load_enriched_questions(enriched_file)
    logger.info(f"Загружено {len(questions)} вопросов из {enriched_file.name}")

    # Фильтр по диапазону номеров
    if args.question_range:
        parts = args.question_range.split("-")
        q_start, q_end = int(parts[0]), int(parts[1])
        questions = [q for q in questions if q_start <= q["number"] <= q_end]
        logger.info(f"Фильтр: вопросы {q_start}–{q_end}, осталось {len(questions)}")

    # Порядок авторов: primary первые, secondary — в конце
    authors_cfg       = config.get("authors", [])
    if not authors_cfg:
        # Fallback если секция отсутствует в config.yaml
        authors_cfg = [
            {"name": "Гайворонский", "primary": True},
            {"name": "Сапин",        "primary": True},
            {"name": "Якимов",       "primary": False},
        ]
    primary_authors   = [a["name"] for a in authors_cfg if a.get("primary", True)]
    secondary_authors = [a["name"] for a in authors_cfg if not a.get("primary", True)]
    all_authors       = primary_authors + secondary_authors
    logger.info(f"Авторы: {', '.join(all_authors)}")

    # ── Embedding модель + ChromaDB ──
    logger.info(f"Загружаю embedding модель ({device})...")
    embed_model = SentenceTransformer(config["embedding"]["model"], device=device)

    client = chromadb.PersistentClient(path=str(ROOT / config["paths"]["chroma_dir"]))
    try:
        collection = client.get_collection("anatomy")
        logger.info(f"Индекс: {collection.count()} чанков")
    except Exception:
        logger.error("Индекс не найден — сначала запустите:  python scripts/2_build_index.py")
        sys.exit(1)

    # ── Чекпойнт ──
    checkpoint: dict[int, dict] = load_checkpoint(output_path) if args.resume else {}
    if checkpoint:
        logger.info(f"Возобновление: {len(checkpoint)} ответов уже готово")

    results: list[dict] = []

    for question in tqdm(questions, desc="Генерирую ответы"):
        q_num = question["number"]

        if q_num in checkpoint:
            results.append(checkpoint[q_num])
            continue

        # Список retrieval-запросов: основной вопрос + подвопросы
        sub_qs  = question.get("sub_questions", [])
        queries = [question["question"]] + [sq["text"] for sq in sub_qs]

        # ── Retrieval по каждому автору ──
        ctx_by_author:    dict[str, str]        = {}
        chunks_by_author: dict[str, list[dict]] = {}

        for author in all_authors:
            chunks = retrieve_for_author(
                queries, author, collection, embed_model,
                config, max_chunks, top_k_per_query, hyde=hyde,
            )
            chunks_by_author[author] = chunks
            ctx_by_author[author]    = format_context(chunks) if chunks else ""

        # ── Генерация ответа ──
        if args.no_llm:
            parts = [
                f"=== {a} ===\n{ctx}" if ctx else f"=== {a} ===\n(нет данных)"
                for a, ctx in ctx_by_author.items()
            ]
            answer = "\n\n".join(parts)
        else:
            has_context = any(chunks_by_author.values())
            if not has_context:
                answer = "Информация не найдена в базе знаний."
            else:
                user_prompt = build_user_prompt(question, ctx_by_author)
                try:
                    answer = call_llm(
                        user_prompt, SYSTEM_PROMPT_V2, config, args.backend, args.model
                    )
                except Exception as e:
                    logger.error(f"LLM ошибка для вопроса №{q_num}: {e}")
                    answer = (
                        f"[Ошибка LLM: {e}]\n\n"
                        + "\n\n".join(
                            f"=== {a} ===\n{ctx}"
                            for a, ctx in ctx_by_author.items()
                            if ctx
                        )
                    )

        entry = {"answer": answer, "chunks_by_author": chunks_by_author}
        results.append(entry)

        if args.resume:
            checkpoint[q_num] = entry
            save_checkpoint(output_path, checkpoint)

    # ── Сборка и сохранение DOCX ──
    logger.info("Собираю итоговый документ...")
    doc = build_output_doc(questions, results)
    doc.save(str(output_path))
    logger.info(f"Сохранено: {output_path}")
    logger.info(f"Всего вопросов: {len(questions)}")


if __name__ == "__main__":
    main()
