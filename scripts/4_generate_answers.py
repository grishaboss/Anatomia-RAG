#!/usr/bin/env python3
"""
Шаг 4: Генерация развёрнутых ответов на все экзаменационные вопросы.

Что делает:
  • Читает вопросы из DOCX (exam_questions_file из config.yaml)
  • Для каждого вопроса ищет релевантные чанки в базе знаний
  • Генерирует ответ через LLM на основе найденного контекста
  • Сохраняет всё в форматированный DOCX

Использование:
    python scripts/4_generate_answers.py
    python scripts/4_generate_answers.py --backend openai --model gpt-4o
    python scripts/4_generate_answers.py --no-llm     # только контекст, без LLM
    python scripts/4_generate_answers.py --resume     # продолжить прерванную генерацию
"""
from __future__ import annotations

import argparse
import json
import logging
import re
import sys
from pathlib import Path

import chromadb
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from sentence_transformers import SentenceTransformer
from tqdm import tqdm

sys.path.insert(0, str(Path(__file__).parent))
from utils import ROOT, call_llm, format_context, get_device, load_config, retrieve

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger(__name__)

# ──────────────────────────────────────────────
#  Системный промпт
# ──────────────────────────────────────────────

SYSTEM_PROMPT = """Ты эксперт-анатом, помогающий студенту-медику подготовиться к экзамену по анатомии человека.

Отвечай СТРОГО на основе предоставленного контекста из учебников (Сапин, методички кафедры).
НЕ придумывай факты, которых нет в контексте.

Структура ответа:
1. Краткое определение / общая характеристика
2. Строение (части, отделы, слои — по анатомической иерархии)
3. Функции
4. Топография и синтопия (если применимо)
5. Кровоснабжение и иннервация (если есть в контексте)
6. Возрастные / клинические особенности (если есть в контексте)

Пиши академическим языком на русском. При первом упоминании анатомического термина давай латинский эквивалент в скобках."""


# ──────────────────────────────────────────────
#  Парсинг вопросов из DOCX
# ──────────────────────────────────────────────

def _md_path_for_docx(docx_path, config):
    processed_dir = ROOT / config["paths"]["processed_dir"]
    md = processed_dir / (docx_path.stem + ".md")
    return md if md.exists() else None


def _parse_questions_from_md(md_path):
    """Parse questions from Docling-produced markdown."""
    questions = []
    current_section = "Общие вопросы"
    global_num = 0
    for raw_line in md_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        # Раздел: **Текст** или ## Заголовок
        section_m = re.match(r"^(?:\*\*|#+)\s*(.+?)\s*(?:\*\*)?$", line)
        if section_m and not re.match(r"^\d+\.", line):
            candidate = section_m.group(1).strip()
            # пропускаем шапку документа (содержит год 20XX)
            if len(candidate) >= 10 and not re.search(r"\d{4}", candidate):
                current_section = candidate
            continue
        # Вопрос: "1. Текст"
        q_m = re.match(r"^(\d+)\.\s+(.+)", line)
        if q_m:
            global_num += 1
            questions.append({
                "number": global_num,
                "text": q_m.group(2).strip(),
                "section": current_section,
            })
    return questions


def _parse_questions_from_docx(docx_path):
    """Fallback: parse questions directly from DOCX."""
    from docx import Document as _DocxDoc
    questions = []
    current_section = "Общие вопросы"
    auto_num = 0
    for para in _DocxDoc(str(docx_path)).paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower()
        is_heading = "heading" in style or "заголов" in style
        is_bold = bool(para.runs) and all(r.bold for r in para.runs if r.text.strip())
        if is_heading or is_bold:
            current_section = text
            continue
        m = re.match(r"^(\d+)[\.)\]]\s+(.+)", text)
        if m:
            auto_num = max(auto_num, int(m.group(1)))
            questions.append({"number": int(m.group(1)), "text": m.group(2).strip(), "section": current_section})
        elif len(text) >= 15:
            auto_num += 1
            questions.append({"number": auto_num, "text": text, "section": current_section})
    return questions


def parse_exam_questions(docx_path, config=None):
    """Читает вопросы из .md (data/processed/), fallback — из DOCX."""
    md_path = _md_path_for_docx(docx_path, config) if config else None
    if md_path:
        return _parse_questions_from_md(md_path)
    logger.warning(f"MD не найден, читаю DOCX: {docx_path.name}")
    return _parse_questions_from_docx(docx_path)


def ckpt_path(output_path: Path) -> Path:
    return output_path.with_suffix(".checkpoint.json")


def load_checkpoint(output_path: Path) -> dict[int, dict]:
    p = ckpt_path(output_path)
    if p.exists():
        with open(p, encoding="utf-8") as f:
            raw = json.load(f)
        return {int(k): v for k, v in raw.items()}
    return {}


def save_checkpoint(output_path: Path, data: dict[int, dict]) -> None:
    with open(ckpt_path(output_path), "w", encoding="utf-8") as f:
        json.dump({str(k): v for k, v in data.items()}, f, ensure_ascii=False, indent=2)


# ──────────────────────────────────────────────
#  Построение выходного DOCX
# ──────────────────────────────────────────────

def _unique_sources(chunks: list[dict]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for chunk in chunks:
        m = chunk["meta"]
        s = m.get("source_name", "")
        page = m.get("page")
        label = f"{s} (стр. {page})" if page else s
        if label not in seen:
            seen.add(label)
            result.append(label)
    return result


def build_output_doc(
    questions: list[dict],
    results: list[dict],
) -> DocxDocument:
    doc = DocxDocument()

    # Заголовок документа
    title = doc.add_heading("Ответы на экзаменационные вопросы\nпо анатомии человека", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    current_section: str | None = None

    for question, res in zip(questions, results):
        # Заголовок раздела
        if question["section"] != current_section:
            current_section = question["section"]
            sec_heading = doc.add_heading(current_section, level=1)
            sec_heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0x76)

        # Вопрос
        q_heading = doc.add_heading(
            f"Вопрос {question['number']}. {question['text']}", level=2
        )
        for run in q_heading.runs:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        # Ответ
        answer_text = res.get("answer", "")
        for line in answer_text.split("\n"):
            line = line.strip()
            if line:
                p = doc.add_paragraph(line)
                p.style = "Normal"
                p.paragraph_format.space_after = Pt(2)

        # Источники
        sources = _unique_sources(res.get("chunks", []))
        if sources:
            src_para = doc.add_paragraph()
            src_para.paragraph_format.space_before = Pt(4)
            run_label = src_para.add_run("Источники: ")
            run_label.bold = True
            run_label.font.size = Pt(9)
            run_src = src_para.add_run("; ".join(sources))
            run_src.italic = True
            run_src.font.size = Pt(9)
            run_src.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        doc.add_paragraph()  # отступ между вопросами

    return doc


# ──────────────────────────────────────────────
#  Точка входа
# ──────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(description="Генерация ответов на экзаменационные вопросы")
    parser.add_argument("--backend", choices=["ollama", "openai"], help="LLM backend")
    parser.add_argument("--model", help="Название модели")
    parser.add_argument("--no-llm", action="store_true", help="Не вызывать LLM — только контекст")
    parser.add_argument("--resume", action="store_true", help="Продолжить прерванную генерацию")
    args = parser.parse_args()

    config = load_config()
    device = get_device(config)

    exam_file   = ROOT / config["output"]["exam_questions_file"]
    output_path = ROOT / config["output"]["answers_file"]
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if not exam_file.exists():
        logger.error(f"Файл с вопросами не найден: {exam_file}")
        sys.exit(1)

    # Парсинг вопросов
    questions = parse_exam_questions(exam_file, config)
    logger.info(f"Загружено {len(questions)} экзаменационных вопросов")
    if not questions:
        logger.error("Вопросы не распознаны")
        sys.exit(1)

    # Embedding + индекс
    logger.info(f"Загружаю embedding модель ({device})...")
    embed_model = SentenceTransformer(config["embedding"]["model"], device=device)

    client = chromadb.PersistentClient(path=str(ROOT / config["paths"]["chroma_dir"]))
    coll_name = config.get("retrieval", {}).get("collection_name", "anatomy")
    try:
        collection = client.get_collection(coll_name)
        logger.info(f"Индекс {coll_name!r}: {collection.count()} чанков")
    except Exception:
        logger.error(f"Индекс {coll_name!r} не найден — запустите 2_build_index.py")
        sys.exit(1)

    # Чекпойнт для --resume
    checkpoint: dict[int, dict] = load_checkpoint(output_path) if args.resume else {}
    if checkpoint:
        logger.info(f"Возобновление: уже сгенерировано {len(checkpoint)} ответов")

    results: list[dict] = []

    for question in tqdm(questions, desc="Генерирую ответы"):
        q_num = question["number"]

        # Восстановить из чекпойнта
        if q_num in checkpoint:
            results.append(checkpoint[q_num])
            continue

        # Поиск релевантного контекста
        chunks = retrieve(question["text"], collection, embed_model, config)

        if args.no_llm or not chunks:
            answer = format_context(chunks) if chunks else "Информация не найдена в базе знаний."
        else:
            context = format_context(chunks)
            prompt = (
                f"Вопрос к экзамену: {question['text']}\n\n"
                f"Контекст из учебников:\n{context}\n\n"
                "Дай развёрнутый, структурированный ответ на экзаменационный вопрос."
            )
            try:
                answer = call_llm(prompt, SYSTEM_PROMPT, config, args.backend, args.model)
            except Exception as e:
                logger.error(f"LLM ошибка для вопроса №{q_num}: {e}")
                answer = f"[Ошибка LLM: {e}]\n\nКонтекст из базы знаний:\n{format_context(chunks)}"

        entry = {"answer": answer, "chunks": chunks}
        results.append(entry)

        # Сохранять чекпойнт после каждого вопроса (актуально при --resume)
        if args.resume:
            checkpoint[q_num] = entry
            save_checkpoint(output_path, checkpoint)

    # Собрать и сохранить DOCX
    logger.info("Собираю итоговый документ...")
    doc = build_output_doc(questions, results)
    doc.save(str(output_path))
    logger.info(f"Сохранено: {output_path}")
    logger.info(f"Всего вопросов: {len(questions)}")


if __name__ == "__main__":
    main()
