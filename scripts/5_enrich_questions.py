#!/usr/bin/env python3
"""
Шаг 5: Семантическое обогащение экзаменационных вопросов подвопросами
из семестровых файлов (questions/).

Что делает:
  • Парсит нумерованные вопросы из главного экзамен-файла (DOCX)
  • Парсит все остальные DOCX в questions/ → flat-список подвопросов
  • Строит embeddings (BAAI/bge-m3) для всех строк
  • Для каждого экзамвопроса: cosine similarity → top-k подвопросов выше threshold
  • Сохраняет data/enriched_questions.json

Использование:
    python scripts/5_enrich_questions.py
    python scripts/5_enrich_questions.py --threshold 0.45 --top-k 20
    python scripts/5_enrich_questions.py --output data/enriched_questions.json
"""
from __future__ import annotations

import argparse
import json
import logging
import re
import sys
from pathlib import Path

import numpy as np
from docx import Document as DocxDocument
from sentence_transformers import SentenceTransformer
from tqdm import tqdm

sys.path.insert(0, str(Path(__file__).parent))
from utils import ROOT, get_device, load_config

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger(__name__)


# ──────────────────────────────────────────────
#  Парсинг DOCX
# ──────────────────────────────────────────────

def _md_path_for_docx(docx_path, config):
    processed_dir = ROOT / config["paths"]["processed_dir"]
    md = processed_dir / (docx_path.stem + ".md")
    return md if md.exists() else None


def _parse_questions_from_md(md_path):
    """Parse questions from Docling-produced markdown."""
    questions = []
    current_section = "Общие вопросы"
    global_num = 0
    for raw_line in md_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        # Раздел: **Текст** или ## Заголовок
        section_m = re.match(r"^(?:\*\*|#+)\s*(.+?)\s*(?:\*\*)?$", line)
        if section_m and not re.match(r"^\d+\.", line):
            candidate = section_m.group(1).strip()
            # пропускаем шапку документа (содержит год 20XX)
            if len(candidate) >= 10 and not re.search(r"\d{4}", candidate):
                current_section = candidate
            continue
        # Вопрос: "1. Текст"
        q_m = re.match(r"^(\d+)\.\s+(.+)", line)
        if q_m:
            global_num += 1
            questions.append({
                "number": global_num,
                "text": q_m.group(2).strip(),
                "section": current_section,
            })
    return questions


def _parse_questions_from_docx(docx_path):
    """Fallback: parse questions directly from DOCX."""
    from docx import Document as _DocxDoc
    questions = []
    current_section = "Общие вопросы"
    auto_num = 0
    for para in _DocxDoc(str(docx_path)).paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower()
        is_heading = "heading" in style or "заголов" in style
        is_bold = bool(para.runs) and all(r.bold for r in para.runs if r.text.strip())
        if is_heading or is_bold:
            current_section = text
            continue
        m = re.match(r"^(\d+)[\.)\]]\s+(.+)", text)
        if m:
            auto_num = max(auto_num, int(m.group(1)))
            questions.append({"number": int(m.group(1)), "text": m.group(2).strip(), "section": current_section})
        elif len(text) >= 15:
            auto_num += 1
            questions.append({"number": auto_num, "text": text, "section": current_section})
    return questions


def parse_exam_questions(docx_path, config=None):
    """Читает вопросы из .md (data/processed/), fallback — из DOCX."""
    md_path = _md_path_for_docx(docx_path, config) if config else None
    if md_path:
        return _parse_questions_from_md(md_path)
    logger.warning(f"MD не найден, читаю DOCX: {docx_path.name}")
    return _parse_questions_from_docx(docx_path)


def parse_sub_questions(docx_path: Path, config: dict | None = None) -> list[dict]:
    """Извлечь подвопросы из семестрового файла.

    Предпочитает .md из data/processed/ (чистый текст Docling).
    Fallback: парсинг DOCX.
    """
    source_name = docx_path.name
    md_path: Path | None = None
    if config is not None:
        md_path = _md_path_for_docx(docx_path, config)

    if md_path:
        return _parse_sub_questions_from_md(md_path, source_name)
    return _parse_sub_questions_from_docx(docx_path, source_name)


def _parse_sub_questions_from_md(md_path: Path, source_name: str) -> list[dict]:
    """Парсинг подвопросов из MD-файла (Docling output).

    Правила:
    - Заголовки (**жирный**, ##) пропускаем
    - Нумерованные строки "1. текст" → берём текст
    - Строки-вопросы (заканчиваются на ?) → берём
    - Остальные непустые строки ≥ 20 символов → берём как подвопросы
    - Строки, которые выглядят как перечни демонстраций (≥ 3 пунктов через \n),
      разбиваем на отдельные элементы
    """
    sub_questions: list[dict] = []
    for raw_line in md_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        # Пропускаем заголовки разделов
        if re.match(r"^(?:\*\*|#+)\s*.+", line) and not re.match(r"^\d+\.", line):
            continue
        # Убираем номер в начале
        clean = re.sub(r"^\d+[\.\)]\s+", "", line).strip()
        if len(clean) < 15:
            continue
        sub_questions.append({"text": clean, "source": source_name})
    return sub_questions


def _parse_sub_questions_from_docx(docx_path: Path, source_name: str) -> list[dict]:
    """Fallback: парсинг подвопросов из DOCX."""
    sub_questions: list[dict] = []
    for para in DocxDocument(str(docx_path)).paragraphs:
        text = para.text.strip()
        if not text or len(text) < 10:
            continue
        style = para.style.name.lower()
        if ("heading" in style or "заголов" in style) and len(text) < 60:
            continue
        is_numbered = bool(re.match(r"^\d+[\.\)]\s", text))
        is_question = text.endswith("?")
        if is_numbered or is_question or len(text) >= 20:
            clean = re.sub(r"^\d+[\.\)]\s+", "", text).strip()
            if clean and len(clean) >= 15:
                sub_questions.append({"text": clean, "source": source_name})
    return sub_questions


def collect_sub_question_files(config: dict, exam_file: Path) -> list[Path]:
    """Найти все DOCX в questions/, исключая главный экзамен-файл."""
    questions_dir = ROOT / config["paths"]["questions_dir"]
    result = []
    for f in sorted(questions_dir.rglob("*.docx")):
        if f.resolve() != exam_file.resolve():
            result.append(f)
    return result


# ──────────────────────────────────────────────
#  Косинусное сходство
# ──────────────────────────────────────────────

def cosine_similarity_matrix(query_emb: np.ndarray, matrix: np.ndarray) -> np.ndarray:
    """
    Вычислить cosine similarity одного вектора query_emb (D,) со матрицей (N, D).
    Оба должны быть L2-нормализованы → dot product = cosine similarity.
    """
    return matrix @ query_emb  # (N,)


# ──────────────────────────────────────────────
#  Точка входа
# ──────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Семантическое обогащение экзаменационных вопросов подвопросами"
    )
    parser.add_argument("--threshold", type=float, help="Минимальное сходство (по умолчанию из config)")
    parser.add_argument("--top-k",     type=int,   help="Макс. подвопросов на вопрос (из config)")
    parser.add_argument("--output",                help="Путь к выходному JSON файлу")
    args = parser.parse_args()

    config  = load_config()
    enrich  = config.get("enrichment", {})
    device  = get_device(config)

    threshold   = args.threshold if args.threshold is not None else float(enrich.get("threshold",   0.50))
    top_k       = args.top_k     if args.top_k     is not None else int(enrich.get("top_k",         15))
    output_file = ROOT / (args.output or enrich.get("output_file", "data/enriched_questions.json"))

    exam_file = ROOT / config["output"]["exam_questions_file"]
    if not exam_file.exists():
        logger.error(f"Экзамен-файл не найден: {exam_file}")
        sys.exit(1)

    # ── Парсинг экзаменационных вопросов ──
    logger.info("Парсинг экзаменационных вопросов...")
    exam_questions = parse_exam_questions(exam_file, config)
    logger.info(f"  Загружено {len(exam_questions)} вопросов")
    if not exam_questions:
        logger.error("Вопросы не распознаны")
        sys.exit(1)

    # ── Парсинг подвопросов из семестровых файлов ──
    sub_q_files = collect_sub_question_files(config, exam_file)
    logger.info(f"Парсинг {len(sub_q_files)} семестровых файлов...")
    all_sub_questions: list[dict] = []
    for f in sub_q_files:
        sqs = parse_sub_questions(f, config)
        logger.info(f"  {f.name}: {len(sqs)} подвопросов")
        all_sub_questions.extend(sqs)

    logger.info(f"Итого подвопросов: {len(all_sub_questions)}")
    if not all_sub_questions:
        logger.error("Подвопросы не найдены ни в одном файле")
        sys.exit(1)

    # ── Загрузка модели ──
    model_name = config["embedding"]["model"]
    batch_size = int(config["embedding"]["batch_size"])
    logger.info(f"Загружаю модель {model_name} ({device})...")
    model = SentenceTransformer(model_name, device=device)

    # ── Embeddings подвопросов ──
    sq_texts = [sq["text"] for sq in all_sub_questions]
    logger.info(f"Вычисляю embeddings для {len(sq_texts)} подвопросов...")
    sq_embeddings = model.encode(
        sq_texts,
        batch_size=batch_size,
        normalize_embeddings=True,
        show_progress_bar=True,
        convert_to_numpy=True,
    )  # shape: (N_sub, D)

    # ── Embeddings экзаменационных вопросов ──
    exam_texts = [q["text"] for q in exam_questions]
    logger.info(f"Вычисляю embeddings для {len(exam_texts)} экзаменационных вопросов...")
    exam_embeddings = model.encode(
        exam_texts,
        batch_size=batch_size,
        normalize_embeddings=True,
        show_progress_bar=True,
        convert_to_numpy=True,
    )  # shape: (N_exam, D)

    # ── Сопоставление ──
    logger.info("Сопоставление подвопросов с экзаменационными вопросами...")
    enriched: list[dict] = []

    for question, q_emb in tqdm(
        zip(exam_questions, exam_embeddings), total=len(exam_questions), desc="Сопоставление"
    ):
        scores = cosine_similarity_matrix(q_emb, sq_embeddings)  # (N_sub,)

        above_idx = np.where(scores >= threshold)[0]
        if len(above_idx) == 0:
            matched: list[dict] = []
        else:
            # Сортируем по убыванию и берём top_k
            sorted_idx = above_idx[np.argsort(scores[above_idx])[::-1]][:top_k]
            matched = [
                {
                    "text":   all_sub_questions[int(j)]["text"],
                    "source": all_sub_questions[int(j)]["source"],
                    "score":  round(float(scores[j]), 4),
                }
                for j in sorted_idx
            ]

        enriched.append(
            {
                "number":        question["number"],
                "question":      question["text"],
                "section":       question["section"],
                "sub_questions": matched,
            }
        )

    # ── Сохранение ──
    output_file.parent.mkdir(parents=True, exist_ok=True)
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(enriched, f, ensure_ascii=False, indent=2)

    total_subs = sum(len(e["sub_questions"]) for e in enriched)
    avg_subs   = total_subs / len(enriched) if enriched else 0.0
    covered    = sum(1 for e in enriched if e["sub_questions"])

    logger.info(f"Сохранено: {output_file}")
    logger.info(
        f"Итого: {len(enriched)} вопросов, "
        f"{total_subs} подвопросов (ср. {avg_subs:.1f}/вопрос), "
        f"{covered}/{len(enriched)} вопросов получили подвопросы"
    )


if __name__ == "__main__":
    main()
